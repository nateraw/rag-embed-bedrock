import json
import pickle
import fire
import boto3
import numpy as np
from io import BytesIO
from pathlib import Path
from typing import List, Dict, Tuple, Optional
from dataclasses import dataclass, asdict
import PyPDF2
from docx import Document
import logging
from datetime import datetime

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


@dataclass
class DocumentChunk:
    """Represents a chunk of document text with metadata."""
    doc_id: str
    chunk_id: int
    text: str
    source_file: str
    page_number: Optional[int] = None
    metadata: Dict = None


class TextSplitter:
    """Simple text splitter that chunks text by character count with overlap."""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def split_text(self, text: str) -> List[str]:
        """Split text into chunks with overlap."""
        chunks = []
        start = 0
        text_length = len(text)
        
        while start < text_length:
            end = start + self.chunk_size
            
            # Try to find a sentence boundary
            if end < text_length:
                # Look for period, question mark, or exclamation point
                for i in range(min(50, end - start), 0, -1):
                    if end - i < text_length and text[end - i] in '.!?':
                        end = end - i + 1
                        break
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - self.chunk_overlap
        
        return chunks


class DocumentProcessor:
    """Processes documents and extracts text."""
    
    @staticmethod
    def extract_text_from_pdf(file_buffer: BytesIO) -> List[Tuple[str, int]]:
        """Extract text from PDF file with page numbers."""
        pdf_reader = PyPDF2.PdfReader(file_buffer)
        pages_text = []
        
        for page_num, page in enumerate(pdf_reader.pages, 1):
            text = page.extract_text()
            if text.strip():
                pages_text.append((text, page_num))
        
        return pages_text
    
    @staticmethod
    def extract_text_from_docx(file_buffer: BytesIO) -> str:
        """Extract text from DOCX file."""
        doc = Document(file_buffer)
        full_text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        return '\n'.join(full_text)


class EmbeddingGenerator:
    """Generates embeddings using Amazon Bedrock with Cohere Embed v3."""
    
    def __init__(self, region_name: str = 'us-east-1'):
        self.bedrock = boto3.client(
            service_name='bedrock-runtime',
            region_name=region_name
        )
        self.model_id = 'cohere.embed-english-v3'
    
    def create_embedding(self, text: str, input_type: str = 'search_document') -> np.ndarray:
        """Create embedding for a single text chunk using Cohere."""
        body = json.dumps({
            "texts": [text],
            "input_type": input_type,
            "embedding_types": ["float"]
        })
        
        response = self.bedrock.invoke_model(
            body=body,
            modelId=self.model_id,
            accept='application/json',
            contentType='application/json'
        )
        
        response_body = json.loads(response['body'].read())
        # Cohere returns embeddings in a different structure
        embedding = np.array(response_body['embeddings']['float'][0])
        return embedding
    
    def create_embeddings_batch(self, texts: List[str], input_type: str = 'search_document') -> np.ndarray:
        """Create embeddings for multiple texts using Cohere."""
        embeddings = []
        # Cohere can handle batches, but has limits. Process in chunks of 96
        batch_size = 96
        
        for i in range(0, len(texts), batch_size):
            batch_texts = texts[i:i + batch_size]
            logger.info(f"Creating embeddings for batch {i//batch_size + 1}/{(len(texts) + batch_size - 1)//batch_size}")
            
            body = json.dumps({
                "texts": batch_texts,
                "input_type": input_type,
                "embedding_types": ["float"]
            })
            
            response = self.bedrock.invoke_model(
                body=body,
                modelId=self.model_id,
                accept='application/json',
                contentType='application/json'
            )
            
            response_body = json.loads(response['body'].read())
            batch_embeddings = response_body['embeddings']['float']
            embeddings.extend(batch_embeddings)
        
        return np.array(embeddings)


class DocumentEmbedder:
    """Main class for processing documents and creating embeddings."""
    
    def __init__(
        self,
        bucket_name: str,
        prefix: str,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        region_name: str = 'us-east-1'
    ):
        self.bucket_name = bucket_name
        self.prefix = prefix
        self.s3 = boto3.client('s3', region_name=region_name)
        self.text_splitter = TextSplitter(chunk_size, chunk_overlap)
        self.embedding_generator = EmbeddingGenerator(region_name)
        self.processor = DocumentProcessor()
    
    def list_documents(self) -> List[str]:
        """List all documents in the S3 bucket with the specified prefix."""
        response = self.s3.list_objects_v2(
            Bucket=self.bucket_name,
            Prefix=self.prefix
        )
        
        documents = []
        for obj in response.get('Contents', []):
            key = obj['Key']
            if key.endswith(('.pdf', '.docx')):
                documents.append(key)
        
        logger.info(f"Found {len(documents)} documents")
        return documents
    
    def process_document(self, s3_key: str) -> List[DocumentChunk]:
        """Process a single document and return chunks."""
        logger.info(f"Processing {s3_key}")
        
        # Download document to memory
        response = self.s3.get_object(Bucket=self.bucket_name, Key=s3_key)
        file_buffer = BytesIO(response['Body'].read())
        
        chunks = []
        doc_id = Path(s3_key).stem
        
        if s3_key.endswith('.pdf'):
            pages_text = self.processor.extract_text_from_pdf(file_buffer)
            for page_text, page_num in pages_text:
                text_chunks = self.text_splitter.split_text(page_text)
                for i, chunk_text in enumerate(text_chunks):
                    chunk = DocumentChunk(
                        doc_id=doc_id,
                        chunk_id=len(chunks),
                        text=chunk_text,
                        source_file=s3_key,
                        page_number=page_num,
                        metadata={'file_type': 'pdf'}
                    )
                    chunks.append(chunk)
        
        elif s3_key.endswith('.docx'):
            full_text = self.processor.extract_text_from_docx(file_buffer)
            text_chunks = self.text_splitter.split_text(full_text)
            for i, chunk_text in enumerate(text_chunks):
                chunk = DocumentChunk(
                    doc_id=doc_id,
                    chunk_id=i,
                    text=chunk_text,
                    source_file=s3_key,
                    metadata={'file_type': 'docx'}
                )
                chunks.append(chunk)
        
        logger.info(f"Created {len(chunks)} chunks from {s3_key}")
        return chunks
    
    def process_all_documents(self) -> Tuple[List[DocumentChunk], np.ndarray]:
        """Process all documents and create embeddings."""
        documents = self.list_documents()
        all_chunks = []
        
        # Process each document
        for doc in documents:
            chunks = self.process_document(doc)
            all_chunks.extend(chunks)
        
        logger.info(f"Total chunks created: {len(all_chunks)}")
        
        # Create embeddings with search_document input type
        texts = [chunk.text for chunk in all_chunks]
        embeddings = self.embedding_generator.create_embeddings_batch(
            texts, 
            input_type='search_document'  # Important for document indexing
        )
        
        return all_chunks, embeddings
    
    def save_embeddings(
        self,
        chunks: List[DocumentChunk],
        embeddings: np.ndarray,
        output_dir: str = './embeddings'
    ):
        """Save chunks and embeddings to local disk."""
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)
        
        # Save chunks as JSON
        chunks_data = [asdict(chunk) for chunk in chunks]
        with open(output_path / 'chunks.json', 'w') as f:
            json.dump(chunks_data, f, indent=2)
        
        # Save embeddings as numpy array
        np.save(output_path / 'embeddings.npy', embeddings)
        
        # Save metadata
        metadata = {
            'created_at': datetime.now().isoformat(),
            'num_chunks': len(chunks),
            'embedding_dim': embeddings.shape[1],
            'embedding_model': self.embedding_generator.model_id,
            'chunk_size': self.text_splitter.chunk_size,
            'chunk_overlap': self.text_splitter.chunk_overlap
        }
        with open(output_path / 'metadata.json', 'w') as f:
            json.dump(metadata, f, indent=2)
        
        logger.info(f"Saved embeddings to {output_path}")
    
    def upload_to_s3(self, local_dir: str, s3_prefix: str):
        """Upload embedding files to S3."""
        local_path = Path(local_dir)
        
        for file_path in local_path.glob('*'):
            if file_path.is_file():
                s3_key = f"{s3_prefix}/{file_path.name}"
                logger.info(f"Uploading {file_path} to s3://{self.bucket_name}/{s3_key}")
                self.s3.upload_file(str(file_path), self.bucket_name, s3_key)


def main(
    bucket_name: str = 'my-bucket',
    prefix: str = 'my_docs/',
    output_dir: str = './embeddings',
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
    upload_to_s3: bool = False,
    s3_output_prefix: str = 'embeddings/',
    region: str = 'us-east-1'
):
    """
    Process documents from S3 and create embeddings for RAG.
    
    Args:
        bucket_name: S3 bucket name
        prefix: S3 prefix for documents
        output_dir: Local directory to save embeddings
        chunk_size: Size of text chunks
        chunk_overlap: Overlap between chunks
        upload_to_s3: Whether to upload results back to S3
        s3_output_prefix: S3 prefix for output files
        region: AWS region
    """
    embedder = DocumentEmbedder(
        bucket_name=bucket_name,
        prefix=prefix,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        region_name=region
    )
    
    # Process documents and create embeddings
    chunks, embeddings = embedder.process_all_documents()
    
    # Save locally
    embedder.save_embeddings(chunks, embeddings, output_dir)
    
    # Optionally upload to S3
    if upload_to_s3:
        embedder.upload_to_s3(output_dir, s3_output_prefix)
        logger.info(f"Uploaded embeddings to s3://{bucket_name}/{s3_output_prefix}")
    
    logger.info("Processing complete!")


if __name__ == '__main__':
    fire.Fire(main)
